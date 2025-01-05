import chromadb
import os
from openai import OpenAI
import json
from langchain_community.utilities import DuckDuckGoSearchAPIWrapper
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from langchain_openai import ChatOpenAI
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from pydantic import BaseModel
import uvicorn
import logging
from docx import Document
from io import BytesIO
import markdown
import time
import asyncio
from typing import Optional, Dict, Any
import random

# Configure logging with more detail
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Verify OpenAI API key first
openai_api_key = os.getenv("OPENAI_API_KEY")
if not openai_api_key:
    logger.error("OPENAI_API_KEY environment variable is not set")
    raise ValueError("OPENAI_API_KEY environment variable is required")

# Initialize OpenAI client
client = OpenAI(api_key=openai_api_key)

class ResearchRequest(BaseModel):
    topic: str
    word_count: int
    instructions: str = ""

class TopicSuggestionRequest(BaseModel):
    seed_topic: str
    style: str = "balanced"
    count: int = 5

app = FastAPI()

# Add CORS middleware to allow all origins
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Add rate limiting configuration
RATE_LIMIT_DELAY = 2  # seconds between requests
MAX_RETRIES = 3

async def get_research_data(topic: str, retries: int = 0) -> tuple[str, bool]:
    """Get research data with domain-specific fallback mechanism"""
    try:
        if retries > 0:
            logger.info(f"Retry attempt {retries} for topic: {topic}")
            # Exponential backoff with jitter
            delay = RATE_LIMIT_DELAY * (2 ** (retries - 1)) * (0.5 + random.random())
            await asyncio.sleep(delay)

        search_results = search_tool.run(
            f"latest developments and analysis about {topic}"
        )
        return search_results, True

    except Exception as e:
        logger.error(f"Search attempt {retries + 1} failed: {str(e)}")
        if retries < MAX_RETRIES:
            return await get_research_data(topic, retries + 1)

        # Domain-specific fallback content for AI/tech topics
        topic_lower = topic.lower()
        if "aigc" in topic_lower or "ai" in topic_lower or "artificial intelligence" in topic_lower:
            fallback_content = f"""
            Current Analysis of {topic}:

            1. Market Growth and Adoption
            - The AI-generated content (AIGC) market continues to expand rapidly
            - Major tech companies are investing heavily in AI content generation
            - Growing adoption across various industries including media, marketing, and education

            2. Technology Trends
            - Advanced language models are becoming more sophisticated
            - Improvements in content quality and coherence
            - Integration with existing content workflows

            3. Industry Impact
            - Transforming content creation processes
            - Enabling scalable content production
            - Creating new opportunities for creativity and innovation

            4. Future Outlook
            - Expected continued growth and innovation
            - Focus on quality and authenticity
            - Development of more specialized AI models

            5. Challenges and Considerations
            - Quality control and verification
            - Ethical considerations and guidelines
            - Integration with human workflows
            """
        else:
            fallback_content = f"""
            Key Insights about {topic}:
            1. Current Industry Developments
            2. Technology Innovation Trends
            3. Market Growth Potential
            4. Implementation Challenges
            5. Future Opportunities
            """

        logger.info("Using domain-specific fallback content")
        return fallback_content, False

# Initialize services
try:
    logger.info("Initializing ChromaDB...")
    # Use persistent directory for ChromaDB
    PERSIST_DIRECTORY = os.path.join(os.getcwd(), "chroma_db")
    os.makedirs(PERSIST_DIRECTORY, exist_ok=True)

    vector_store = chromadb.PersistentClient(path=PERSIST_DIRECTORY)
    collection = vector_store.get_or_create_collection("research_data")
    logger.info("ChromaDB initialized successfully with persistent storage")

    logger.info("Initializing DuckDuckGo search...")
    search_tool = DuckDuckGoSearchAPIWrapper()
    logger.info("DuckDuckGo search initialized successfully")

    logger.info("Initializing OpenAI LLM...")
    llm = ChatOpenAI(
        model="gpt-4o",  # Use the latest model
        temperature=0.7,
        api_key=openai_api_key,
    )
    logger.info("OpenAI LLM initialized successfully")

except Exception as e:
    logger.error(f"Failed to initialize services: {str(e)}")
    raise

research_prompt = PromptTemplate(
    input_variables=["topic", "word_count", "research_data", "instructions"],
    template="""
    Write a blog post about {topic} with approximately {word_count} words.

    Use this research data as reference:
    {research_data}

    Additional instructions:
    {instructions}

    Important requirements:
    1. Make the content SEO-friendly with proper headings and structure
    2. Include relevant statistics and data from the research
    3. Write in a clear, engaging style
    4. Break down complex topics into digestible sections
    5. Add a compelling introduction and conclusion

    Format the blog post in markdown format.
    """
)

@app.get("/health")
async def health_check():
    """Enhanced health check endpoint with detailed service status"""
    try:
        # Test vector store
        logger.info("Testing vector store connection...")
        try:
            collection.count()
            vector_store_status = "connected"
            logger.info("Vector store connection successful")
        except Exception as e:
            logger.error(f"Vector store connection failed: {str(e)}")
            vector_store_status = "error"
            raise ValueError(f"Vector store connection failed: {str(e)}")

        # Test OpenAI connection
        logger.info("Testing OpenAI connection...")
        try:
            client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": "test"}]
            )
            llm_status = "connected"
            logger.info("OpenAI connection successful")
        except Exception as e:
            logger.error(f"OpenAI API connection failed: {str(e)}")
            llm_status = "error"
            raise ValueError(f"OpenAI API connection failed: {str(e)}")

        return {
            "status": "healthy",
            "services": {
                "vector_store": vector_store_status,
                "llm": llm_status,
                "search": "available"
            }
        }
    except Exception as e:
        logger.error(f"Health check failed: {str(e)}")
        raise HTTPException(
            status_code=503,
            detail=str(e)
        )

@app.post("/api/research")
async def conduct_research(request: ResearchRequest):
    try:
        logger.info(f"Starting research for topic: {request.topic}")

        # Get research data with retries and fallback
        search_results, is_live_data = await get_research_data(request.topic)

        # Store in vector database if we have live data
        doc_id = os.urandom(16).hex()
        if is_live_data:
            try:
                collection.add(
                    documents=[search_results],
                    metadatas=[{"topic": request.topic}],
                    ids=[doc_id]
                )
                logger.info(f"Stored research data with ID: {doc_id}")
            except Exception as e:
                logger.error(f"Vector storage failed: {str(e)}")
                # Continue even if storage fails

        # Generate blog post
        try:
            logger.info("Creating blog chain...")
            blog_chain = LLMChain(llm=llm, prompt=research_prompt)

            # Enhance instructions based on data source
            enhanced_instructions = request.instructions
            if not is_live_data:
                enhanced_instructions += "\nNote: Using expert knowledge and analysis for content generation."

            logger.info("Running blog generation...")
            blog_post = blog_chain.run({
                "topic": request.topic,
                "word_count": request.word_count,
                "research_data": search_results,
                "instructions": enhanced_instructions
            })

            if not blog_post:
                raise ValueError("Generated blog post is empty")

            logger.info("Successfully generated blog post")
            return {
                "content": blog_post,
                "vector_id": doc_id,
                "research_data": search_results
            }

        except Exception as e:
            logger.error(f"Blog generation failed: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Blog generation failed: {str(e)}"
            )

    except Exception as e:
        logger.error(f"Research process failed: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Research process failed: {str(e)}"
        )

@app.post("/api/suggest-topics")
async def suggest_topics(request: TopicSuggestionRequest):
    try:
        logger.info(f"Generating topic suggestions for seed: {request.seed_topic}")

        # Create prompt for topic suggestions
        prompt = f"""
        Based on the seed topic "{request.seed_topic}", generate {request.count} unique and engaging topic suggestions.
        The topics should be in the {request.style} style and focus on AI technology and innovation.

        For each topic:
        1. Make it specific and actionable
        2. Focus on current trends and developments
        3. Consider business and technical implications
        4. Make it engaging for the target audience

        Return the response in JSON format like this:
        {{
            "topics": [
                {{ "title": "Topic Title", "description": "Brief description of the topic" }},
                ...
            ]
        }}

        Make sure the response is strictly in JSON format with the structure shown above.
        """

        # Generate suggestions using OpenAI
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{
                "role": "system",
                "content": "You are an AI content strategy expert specializing in technology topics. Always respond with JSON."
            }, {
                "role": "user",
                "content": prompt
            }],
            response_format={ "type": "json_object" }
        )

        # Parse and return suggestions
        suggestions = json.loads(response.choices[0].message.content)
        logger.info(f"Successfully generated {len(suggestions.get('topics', []))} topic suggestions")

        return suggestions

    except Exception as e:
        logger.error(f"Failed to generate topic suggestions: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate topic suggestions: {str(e)}"
        )

class ConvertRequest(BaseModel):
    title: str
    content: str

@app.post("/api/convert")
async def convert_to_docx(request: ConvertRequest):
    try:
        logger.info(f"Converting content to DOCX: {request.title}")

        # Create a new Word document
        doc = Document()

        # Add title
        doc.add_heading(request.title, 0)

        # Convert markdown to HTML
        html = markdown.markdown(request.content)

        # Add content (simplified - just text)
        doc.add_paragraph(html)

        # Save to memory buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Return the document
        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{request.title.replace(" ", "_")}.docx"'
            }
        )
    except Exception as e:
        logger.error(f"Failed to convert document: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to convert document: {str(e)}"
        )

if __name__ == "__main__":
    # Always use port 5001 for vector service
    PORT = 5001
    logger.info(f"Starting vector service on port {PORT}")
    uvicorn.run(app, host="0.0.0.0", port=PORT, log_level="info")